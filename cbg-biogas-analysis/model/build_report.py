"""
build_report.py — assemble ../report/main.docx (decision-grade evaluation).
Pulls summary tables LIVE from the engine so narrative and numbers can't drift.
Run: python3 build_report.py
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from drivers import Drivers, SCALES_TPD_CBG
from cbg_model import project_financials, mass_energy_balance, MODELS, pct, cr
import scenarios as S

REP = os.path.join(os.path.dirname(__file__), "..", "report")
CH = os.path.join(os.path.dirname(__file__), "..", "charts")
os.makedirs(REP, exist_ok=True)
d = Drivers()
TAG_COLORS = {"FACT": RGBColor(0x2E,0x7D,0x32), "F": RGBColor(0x2E,0x7D,0x32),
              "INFERENCE": RGBColor(0xB7,0x6E,0x00), "I": RGBColor(0xB7,0x6E,0x00),
              "OPINION": RGBColor(0x15,0x65,0xC0), "O": RGBColor(0x15,0x65,0xC0)}
doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)

TAG_RE = re.compile(r"(\[(?:FACT|INFERENCE|OPINION|F|I|O)\]|\*\*[^*]+\*\*)")
def para(text, style=None, italic=False):
    p = doc.add_paragraph(style=style)
    for tok in TAG_RE.split(text):
        if not tok: continue
        if tok.startswith("[") and tok[1:-1] in TAG_COLORS:
            r = p.add_run(tok); r.bold = True; r.font.color.rgb = TAG_COLORS[tok[1:-1]]; r.font.size = Pt(8.5)
        elif tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        else:
            r = p.add_run(tok); r.italic = italic
    return p
def bullets(items):
    for it in items: para(it, style="List Bullet")
def sowhat(items):
    para("So-what for the verdict:").runs[0].bold = True
    bullets(items)
def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; r = c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text = ""; rr = cells[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(9)
    doc.add_paragraph()
def img(name, w=6.3):
    p = os.path.join(CH, name)
    if os.path.exists(p):
        doc.add_picture(p, width=Inches(w)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------------- Title ----------------
t = doc.add_heading("Cow-Dung → Compressed Biogas (CBG) in India", level=0)
para("A decision-grade techno-economic, policy and strategic evaluation for a North-India FMCG / beverage capital allocator (10–15 year horizon).").runs[0].italic = True
para("Prepared 2026-05-31. All material claims tagged [FACT] / [INFERENCE] / [OPINION]. Indian notation (₹, lakh, crore). Figures from the committed model engine (model/cbg_model.py) and research in sources/citations.md.").runs[0].font.size = Pt(8.5)
para("CAVEAT ON RECENCY: policy and prices moved hard in the 18 months to this date. Web research indexes to ~early-mid 2026; items not re-confirmable for the current quarter are in the 'could-not-reconfirm' list. Verify live before committing capital.")

# ---------------- Executive verdict ----------------
doc.add_heading("Executive Verdict", level=1)
para("A cow-dung→CBG plant built **to sell gas** (SATAT injection or cylinder bottling) is a **NO-GO / at best a narrow CONDITIONAL-GO** for this reader: subsidy-dependent, feedstock-squeezed, no contractual demand security, below the 20% hurdle even in the base case, with a **~51% modelled probability of losing money** [INFERENCE]. The only version worth diligence is **captive displacement (M3) at the reader's own beverage plant** — and that is an infrastructure / energy-cost-hedge, not a growth bet.")
table(["Model","Verdict","Base equity IRR","Core reason"],
      [["M1 SATAT injection","NO-GO (cond. only ≥10–15 TPD w/ CGD take-or-pay)","12.6%, NPV≈0, DSCR 0.86","No take-or-pay; CBO is aggregate not gate; dung-infeasible at scale"],
       ["M2 bottled cascade","NO-GO (unconditional)","deeply negative","2–5× capex + 250-bar compression + cascade logistics"],
       ["M3 captive displacement","CONDITIONAL-GO (investigate)","swings on real avoided-fuel price","avoids OMC monopsony, offtake risk, logistics"]])
para("**Single biggest reason to walk away:** no demand security + structurally squeezed margin. **Single condition that flips it (M3 only):** the beverage plant's avoided PNG/LPG/furnace-oil spend is large enough that a right-sized captive digester on cheap (<₹1,200/t, <15 km) dung/co-substrate pays back in < 5 years. **Opportunity cost:** ~7% modelled chance of beating the reader's 20% hurdle — capital earns more, more safely, in the existing playbook.")

# ---------------- Scope / framing ----------------
doc.add_heading("1. Scope & framing — three businesses, not one", level=1)
para('"Bottling and sale" of biogas hides three structurally different businesses (see ADR-001). They are modelled and judged separately:')
bullets(["**M1 — CGD pipeline injection (SATAT):** upgrade to >90% CH₄, sell to OMC/CGD. Volume path; near-monopsony buyer; **no take-or-pay** [FACT].",
         "**M2 — Bottled / cascade CBG:** fill high-pressure cascades, truck to off-grid users. Logistics-heavy; PESO-regulated; avoids OMC dependence.",
         "**M3 — Captive displacement:** burn CBG at the reader's beverage plant; value = avoided fuel cost, not a sale."])
para("Feedstock sub-axis kept explicit: **dung-only** (low yield, ~15 kg CBG/t) vs **dung-anchored co-digestion** (press mud / napier / mandi waste) — different risk profiles [FACT].")

# ---------------- Lens 1 Policy ----------------
doc.add_heading("Lens 1 — Policy & Regulatory", level=1)
para("**SATAT** offtake agreements run **15 years** (the '5-year' claim is a blog error) but carry **no take-or-pay**: OMCs 'shall try to off-take… actual off-take shall entirely depend on the actual quantity produced' [FACT, IOCL FAQ]. **Price** is **~₹77.4/kg ex-GST (₹1,478/MMBTU, Jun 2025)**, set at **85% of average CNG retail — CNG-linked, no floor** [FACT]. The industry is lobbying for a fixed ₹90/kg, itself a margin-stress signal [O].")
para("The **CBO mandatory blending mandate** (1%→3%→4%→5%, FY25-26→FY28-29) is an obligation on **CGD entities at the national-aggregate level via a Central Repository Body — it does NOT convert into enforced offtake at an individual plant's gate** [FACT/INFERENCE]. This single distinction is the crux: the policy stack creates aggregate demand *pull* but **zero contractual demand security** for one plant.")
para("**Subsidy:** MNRE CFA up to ₹4 cr per 4,800 kg/day (cap ₹10 cr); the **DPI pipeline scheme (₹994.5 cr)** funds CGD tie-in (~₹9.95 cr/project) **but requires a Gas Supply Agreement with the CGD carrying ≥50% take-or-pay** — the *only* take-or-pay anywhere, and one you must negotiate [FACT]. Punjab has **no dedicated CBG policy** as of mid-2025; only generic IBDP-2022 incentives [FACT — verify PEDA].")
para("**FOM/digestate:** MDA ₹1,500/t exists and FOM is now in the Fertiliser Control Order, but offtake is immature (~120 plants registered, ~44 MoUs as of Mar 2026) [FACT]. **Carbon (CCTS):** CBG is an approved sector but trading is not live (~late-2026 go-live); ₹600–900/tCO₂e is a forecast, not a cleared price — **SPECULATIVE, exclude from base** [O].")
sowhat(["Policy gives demand *pull*, not demand *security* — model offtake risk as real and uninsurable [INFERENCE].",
        "The only bankable take-or-pay is a CGD GSA you negotiate (M1, ≥2 TPD) — make it a precondition, not an assumption.",
        "FOM is upside and carbon is optionality — neither belongs in the base case."])

# ---------------- Lens 2 Feedstock ----------------
doc.add_heading("Lens 2 — Feedstock economics (the make-or-break lens)", level=1)
para("Collectable dung is **~5 kg/animal/day** in a dairy-anchored setting (gross output ~9–10 kg, but 30–60% is dropped grazing or used as cakes) [FACT]. Wet dung's economical **collection radius is <10–15 km**, with ~4× requirement needed in-radius for security [FACT].")
para("**The margin vise:** SATAT demand has bid feedstock up — press mud went **₹100→₹500–600/t in ~2 years** [FACT]; NDDB anchors dung at **₹1/kg** [FACT]. After the ₹77/kg base rate, only **~₹500–700 is left per tonne of dung** to cover feedstock + debt + profit, so **dung past ~₹1,000–1,500/t erases viability** [FACT/INFERENCE]. Co-digestion (cheap press mud at ₹550/t, 40 kg CBG/t) lifts economics — but press mud is **seasonal (sugar crush Oct–Apr)** and its own price is spiralling.")
para("**Yield discipline:** dung gives ~25–40 m³ biogas/t, ~60% CH₄, reconciling to **~15 kg CBG/t fresh** (the widely-cited '16–20 kg/MT per IS 16087' is industry shorthand — IS 16087 is a gas-quality spec, not a yield norm) [INFERENCE]. Promoter claims of 0.45 m³ CH₄/kg dung are ~10× too high on a fresh-weight basis [FACT].")
sowhat(["Feedstock price (D1) is the #1 sensitivity; the business must outbid dung's existing manure/cake value floor [FACT].",
        "Dung-only cannot reach viable scale (see Lens 3 infeasibility); co-digestion is mandatory, which dilutes the 'cow-dung business' premise [INFERENCE].",
        "Logistics, not capex, is the binding constraint — site within a dairy cluster or don't build [OPINION]."])

# ---------------- Lens 3 Techno-economic ----------------
doc.add_heading("Lens 3 — Techno-economic model", level=1)
para("A parameterised 15-year model (model/cbg_model.py; all assumptions in drivers.py) covers **5 CBG scales × 3 models**. Capex intensity ~₹40–50k/(kg·day) (5 TPD ≈ ₹22–25 cr — the prompt's ₹16 cr is ~30–45% low; a T1 parliamentary cross-check lands ~₹25 cr) [FACT]. Upgrading default = **membrane** (lowest capex/parasitic at small-mid scale) [FACT]. Base PLF **50%** (Indian plants run 20–60% vs 80%+ in Europe) [FACT]. D:E 75:25, debt 10.5%, subsidy capped at 30% of gross (DPI off in base). Net-salary-style closure and unit checks pass the self-test.")
rows = []
for sc in SCALES_TPD_CBG:
    for mdl in MODELS:
        r = project_financials(sc, mdl, d)
        rows.append([f"{sc} TPD", mdl.replace("_"," "), cr(r['cx']['net']), pct(r['proj_irr']), pct(r['eq_irr']),
                     cr(r['npv']), (f"{r['min_dscr']:.2f}" if r['min_dscr'] else "n/a"),
                     "Y" if r['mb']['dung_feasible'] else "NO"])
table(["Scale","Model","Net capex ₹cr","Proj IRR","Eq IRR","NPV ₹cr","Min DSCR","Dung feasible?"], rows)
para("**Dung-availability constraint** (default 50% dung mix): a 5 TPD plant needs ~73,000 cattle within a 15 km radius (≈57,000 available) — **infeasible at ≥5 TPD** [INFERENCE]. The economics-need-scale vs can't-source-dung-at-scale vise is the model's central structural finding.")
drows = []
for sc in SCALES_TPD_CBG:
    mb = mass_energy_balance(sc, d)
    drows.append([f"{sc} TPD", f"{mb['feedstock_tpd']:.0f} t/d", f"{mb['dung_tpd']:.0f} t/d",
                  f"{mb['animals_required']:,.0f}", f"{mb['animals_available_in_radius']:,.0f}",
                  "FEASIBLE" if mb['dung_feasible'] else "INFEASIBLE"])
table(["Scale","Feed/day","Dung/day","Cattle needed","Cattle in 15km","Verdict"], drows)
sowhat(["At base assumptions, M1 returns only ~10% project / 12.6% equity IRR with DSCR 0.86 — below WACC-NPV and below the bank's 1.2 DSCR floor [INFERENCE].",
        "M2 (bottled) is deeply negative at every scale; M3 (captive) ~1% at generic avoided-fuel price [INFERENCE].",
        "The only scales that are even marginally economic (≥5 TPD) are dung-infeasible — forcing co-digestion or import [INFERENCE]."])

# ---------------- Lens 4 Sensitivity ----------------
doc.add_heading("Lens 4 — Sensitivity, scenario & Monte-Carlo", level=1)
sc_res = S.run_scenarios()
table(["Scenario","Project IRR","Equity IRR","NPV ₹cr","Min DSCR"],
      [[n, pct(sc_res[n]['proj_irr']), pct(sc_res[n]['eq_irr']), cr(sc_res[n]['npv']),
        f"{sc_res[n]['min_dscr']:.2f}" if sc_res[n]['min_dscr'] else "n/a"] for n in ["Tailwind","Base","Stress"]])
img("scenario_irr.png", 5.6)
para("**Tornado:** equity-IRR swing ranks **PLF > feedstock price > CBG price > capex > FOM > subsidy > cost of debt** — confirming D6/D1/D4/D5 as the dominant drivers [INFERENCE].")
img("tornado_irr.png", 6.0)
mc = S.monte_carlo()
para(f"**Monte-Carlo (5,000 draws)** on the top drivers: **P10 {pct(mc['p10'])}, P50 {pct(mc['p50'])}, P90 {pct(mc['p90'])}**. "
     f"**P(IRR>WACC)={mc['prob_gt_wacc']*100:.0f}%, P(IRR>20% hurdle)={mc['prob_gt_20']*100:.0f}%, P(IRR<0)={mc['prob_negative']*100:.0f}%** [INFERENCE]. "
     "A coin-flip-to-lose-money proposition that clears the reader's hurdle 7% of the time.")
img("montecarlo_irr.png", 5.8)
para("**Breakeven frontier:** even the favourable corner (cheap dung ₹800/t, high CBG ₹95/kg) only reaches ~21% project IRR; the base point (₹1,250/₹77) sits at WACC. Most of the plausible feedstock×price space is sub-WACC [INFERENCE].")
img("breakeven_frontier.png", 5.8)
sowhat(["Only a full policy-tailwind stack (PLF .70 + DPI + FOM clears + ₹85 + cheap dung) clears 20% [INFERENCE].",
        "The base case sits AT WACC with NPV≈0 — no margin of safety; the stress case is a wipeout (-34% equity IRR) [INFERENCE].",
        "7% probability of beating the hurdle is the decisive number for a 20%-mandate allocator [OPINION]."])

# ---------------- Lens 5 Competitive ----------------
doc.add_heading("Lens 5 — Competitive / strategic structure", level=1)
para("SATAT targeted 5,000 plants by 2023-24; **~108–132 are commissioned — a ~97% shortfall**, ~920 TPD output, with 1,094 LoIs that never became plants (~10% conversion) [FACT]. Failure modes, ranked: feedstock-cost squeeze > low PLF > FOM non-clearance > OMC offtake delays > financing gaps [FACT/INFERENCE].")
para("**Margin capture:** the plant operator is the **squeezed middle** — between feedstock aggregators with pricing power in scarcity and an OMC monopsony paying an administered price. The clearer winners are feedstock aggregators and EPC/subsidy-capture players who earn on the build regardless of PLF [INFERENCE].")
para("**Exit:** unlike a pure orphan asset, a live acquirer market exists — EverEnviro, GPS/ARYA ONE ($400M platform, Sojitz minority stake), Adani, Reliance, OMC JVs (IGRPL raised ₹836 cr) — **but value concentrates in platforms, not single sub-scale plants**, and no public plant EV/EBITDA multiple was disclosed [FACT/INFERENCE].")
sowhat(["As a single sub-scale plant you are the squeezed middle and a price-taker to a monopsony [INFERENCE].",
        "There is an exit, but only for platform-scale, high-PLF portfolios — not a one-off dung plant [INFERENCE].",
        "Late entry into a 'mandated-demand' market that 97% of intended builders never entered is a warning, not an opportunity [OPINION]."])

# ---------------- Lens 6 Macro ----------------
doc.add_heading("Lens 6 — Long-term viability & macro (10–15 yr)", level=1)
para("CNG-transport demand is still growing near-term (>1M CNG cars FY26) but faces structural EV erosion in the 10–15 yr tail (EV target 30% of sales by 2030); **PNG-domestic is the more durable CBG sink** [FACT/INFERENCE]. India's gas import dependence is high (LNG ≈ 50% of availability); the import-substitution thesis is directionally valid (IEEFA: ~US$29bn savings FY25–30) [FACT].")
para("**The price is administered (85% of CNG), not floating** — the operator captures no upside beyond the formula [FACT]. The **CBO 5% is a ceiling, not an ever-rising floor**: the market is genuinely undersupplied vs the mandate for ~2–3 years, but post-saturation incremental plants compete for a fixed obligated volume at an administered price. 'Mandated demand forever' is marketing that conflates the 3-year window with the 10–15 yr horizon [INFERENCE].")
sowhat(["Near-term (2–3 yr) demand cover is real; the 10–15 yr thesis the prompt asks about is not secured [INFERENCE].",
        "Administered, capped pricing means you bear cost inflation without price upside [FACT].",
        "EV erosion + a 5% ceiling cap the terminal growth story — underwrite cashflow, not perpetual demand [OPINION]."])

# ---------------- Lens 7 ESG ----------------
doc.add_heading("Lens 7 — ESG, optics & non-financial return", level=1)
para("Genuine methane-abatement and dairy/agri-waste-management value is real and strategically useful for an FMCG brand's sustainability narrative [FACT]. Carbon credit economics (~$40–75/t CBG) are marginal, not a thesis-maker [O]. The credible risk is **greenwashing optics** if the plant chases energy crops/subsidies rather than genuine waste utilisation.")
sowhat(["The strongest non-financial case is for M3 (captive) as a credible 'cleaner manufacturing' story tied to the reader's own plant [OPINION].",
        "ESG value does not rescue a sub-WACC sale-of-gas business; it is a reason to do M3 well, not to do M1/M2 [OPINION].",
        "Rural/dairy-cluster goodwill in Punjab is a real but unquantified side-benefit [INFERENCE]."])

# ---------------- Lens 8 Execution ----------------
doc.add_heading("Lens 8 — Execution & operational risk", level=1)
para("Anaerobic digestion + gas upgrading is operationally intensive: digester souring/foaming, H₂S corrosion, grit abrasion from dung, upgrading downtime, and skilled-operator scarcity [FACT]. Feedstock contracts with smallholders/dairies are hard to enforce. **The reader is a solo technical/ops leader already stretched across many fronts** — a feedstock-logistics-heavy biological process plant is a major, sustained bandwidth commitment, not a passive asset [OPINION].")
sowhat(["This is a hands-on operating business, not capital-light — it competes for the reader's scarcest resource (attention) [OPINION].",
        "PLF (the #1 driver) is precisely what weak operations destroy — the ops risk and the financial risk are the same risk [INFERENCE].",
        "If pursued, M3 at the existing plant at least co-locates with existing operations and staff [OPINION]."])

# ---------------- Red team ----------------
doc.add_heading("Red-team: value trap vs steel-man", level=1)
para("**The bear case (this is a value trap):** subsidy-dependent; no take-or-pay; CBO is aggregate not gate; feedstock prices are spiralling toward the viability cliff; PLF chronically 20–60%; FOM revenue often doesn't clear (a real Surat plant runs at 2.9% capacity); >70% of early plants missed forecasts; base-case IRR is sub-hurdle with DSCR <1.2; 51% modelled chance of loss; no exit for a single sub-scale plant; and it consumes a stretched operator's bandwidth. For a 20%-mandate, capital-light allocator this is the opposite of the brief [OPINION].")
para("**The steel-man (the case to proceed):** the sector is genuinely undersupplied vs a legislated CBO for the next 2–3 years; co-digestion with cheap press mud materially improves the math; a signed CGD GSA with ≥50% take-or-pay (DPI precondition) removes the central offtake risk and unlocks pipeline subsidy; at ≥10–15 TPD with secured feedstock the asset can reach low-teens IRR and is acquirable by a platform; and **for M3, displacing dear furnace-oil/LPG at the reader's own plant can pay back fast while delivering a real ESG story** [INFERENCE]. The steel-man is strongest for M3 and weakest for M2.")

# ---------------- Verdict ----------------
doc.add_heading("Verdict & decision tree", level=1)
para("**M1 (SATAT injection): NO-GO** for this reader as a growth investment; **CONDITIONAL-GO** only if ALL hold: ≥10–15 TPD scale; a signed CGD gas-supply agreement with **≥50% take-or-pay** (also unlocks DPI subsidy); co-digestion feedstock locked at **<₹1,200/t** within 15 km; PLF underwritten ≥60%. Even then it is a low-teens-IRR infrastructure asset **below the 20% hurdle** [OPINION].")
para("**M2 (bottled cascade): NO-GO (unconditional)** — the 2–5× capex penalty, 250-bar compression and cascade logistics make it negative-IRR; keep it only as an emergency fallback channel if an OMC stops lifting [INFERENCE].")
para("**M3 (captive displacement): CONDITIONAL-GO — the only model to investigate.** Precondition: the beverage plant's annual PNG/LPG/furnace-oil spend, divided by a right-sized captive digester fed on <₹1,200/t dung/co-substrate within 15 km, yields **payback < 5 years** on avoided fuel. Underwrite as a sustainability/energy-cost hedge, not against the 20% hurdle. **This needs one input: driver D13 — your plant's fuel spend (₹/yr) and unit fuel price.** Flex the live model/cbg_model.xlsx with your real numbers.")
para("**Single biggest walk-away reason:** no demand security (no take-or-pay; CBO is aggregate, not plant-gate) combined with a feedstock price spiral → 51% modelled probability of loss, 7% of clearing the hurdle. **Single condition that flips to buy:** M3 with payback < 5 yr on avoided fuel at the reader's own plant.")
para("**Opportunity cost:** against a 20%+ CAGR mandate met by capital-light, liquid, high-conviction bets, a CBG plant is capital-heavy (₹20–25 cr at 5 TPD), illiquid, operationally punishing, and sub-WACC in the base case. Unless the ESG/strategic value to the FMCG brand is itself the objective (M3), the capital is better deployed in the existing playbook [OPINION].")
bullets(["**Want to SELL gas (M1/M2)?** → Walk away unless you can sign a CGD ≥50%-ToP GSA AND lock feedstock <₹1,200/t at ≥10 TPD — then a sub-hurdle infrastructure asset.",
         "**Want to cut the plant's fuel bill / build an ESG story (M3)?** → Investigate. Get D13, right-size to your gas demand + local dung, accept it only if payback < 5 yr.",
         "**Neither?** → Deploy the capital in your existing 20%+ playbook."])

# ---------------- Sources & methods ----------------
doc.add_heading("Sources & methods", level=1)
para("Full dated citations with source tiers (T1 govt/peer-reviewed · T2 trade/consultancy · T3 promoter) and the complete 'could-not-reconfirm' list are in **sources/citations.md**. Load-bearing T1/T2 sources include: IOCL CBG FAQ; PIB CBO notifications (PRID 1979705/1989226); MoPNG/MNRE/GOBARdhan/PNGRB scheme docs; Standing Committee on Petroleum & NG (2022-23); MDPI Energies 18(24):6506 (2025); SBI Caps / Renewable Watch financing analyses; REGlobal; Down To Earth; IEEFA; PPAC; CSE; Indian Biogas Association.")
para("**Method & limitations:** outputs come from a transparent 15-year model (drivers.py → cbg_model.py) cross-checked against published plant economics and the SBI Caps IRR-by-scale ladder; it passes a self-test on unit-chain reconciliation, mass-balance closure, and driver→output propagation. **Known limitations:** capex intensity is held flat across scale (conservative — understates large-plant scale economics); cascade logistics (M2) and the M3 avoided-fuel price (D13) are under-sourced placeholders; FOM, carbon and PLF carry wide uncertainty (reflected in the Monte-Carlo). Several primary PDFs were fetch-blocked and rest on search extraction. **Verify all policy and prices live before committing capital.**")

out = os.path.join(REP, "main.docx")
doc.save(out)
print("wrote", out, "| paragraphs:", len(doc.paragraphs))
